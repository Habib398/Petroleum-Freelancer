"""End-to-end smoke test for the DOCX admin routes.

Spins up the real Flask app pointed at a throw-away SQLite database and
exercises the full template lifecycle through the test_client:

  1. Login as the seeded admin.
  2. Seed a test station + station_profile (so auto-resolution has data).
  3. Upload a sample .docx template (multipart) and verify fields are parsed.
  4. Update field labels and the ``is_required`` flag.
  5. Publish the template.
  6. Generate a document with manual_values + station_id.
  7. List generated documents and verify the entry is present.
  8. Download the generated .docx and verify the placeholders were
     actually replaced.
  9. Approve the generated document, then verify the status transition.
 10. Generate a second document, cancel it with reason, then verify a
     second cancellation attempt is rejected (terminal state).

Run from the project root::

    .venv/Scripts/python.exe scripts/smoke_docx_routes.py

Exits 0 on success, 1 on any failure.
"""

from __future__ import annotations

import io
import os
import sys
import tempfile
import zipfile
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))


FAILED: list[str] = []


def check(label: str, condition: bool, detail: str = "") -> None:
    if condition:
        print(f"  OK  {label}")
    else:
        FAILED.append(label)
        print(f"  FAIL  {label}")
        if detail:
            print(f"        {detail}")


def make_sample_docx_bytes() -> bytes:
    """Return an in-memory .docx with a representative set of placeholders."""
    from docx import Document
    doc = Document()
    doc.add_heading("Bitacora SASISOPA", level=1)
    doc.add_paragraph("Estacion: <<NOMBRE_ESTACION>>")
    doc.add_paragraph("RFC: <<RFC>>")
    doc.add_paragraph("Permiso CRE: <<PERMISO_CRE>>")
    doc.add_paragraph("Fecha: <<FECHA_HOY>>")
    doc.add_paragraph("Observaciones: <<OBSERVACIONES>>")
    doc.add_paragraph("Hallazgos: <<HALLAZGOS>>")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def docx_text_from_bytes(data: bytes) -> str:
    """Extract concatenated paragraph text from a .docx blob (without using
    python-docx, to keep the verification independent of the engine)."""
    import re
    with zipfile.ZipFile(io.BytesIO(data)) as zf:
        xml = zf.read("word/document.xml").decode("utf-8", errors="ignore")
    # Drop tags, keep text only.
    no_tags = re.sub(r"<[^>]+>", " ", xml)
    return re.sub(r"\s+", " ", no_tags).strip()


def main() -> int:
    print("DOCX routes smoke test")
    print("======================")

    with tempfile.TemporaryDirectory(prefix="docx_routes_smoke_") as tmp:
        tmpdir = Path(tmp)
        db_path = tmpdir / "smoke.db"
        upload_dir = tmpdir / "uploads"
        upload_dir.mkdir(parents=True, exist_ok=True)

        # Pin the app to our throw-away locations BEFORE create_app loads.
        os.environ["COG_DB_PATH"] = str(db_path)
        os.environ["COG_UPLOAD_DIR"] = str(upload_dir)
        os.environ["COG_CSRF"] = "0"          # bypass CSRF (test-only)
        os.environ["COG_SECRET"] = "smoke-test-secret"
        os.environ["COG_ADMIN_USER"] = "admin"
        os.environ["COG_ADMIN_PASS"] = "admin123"
        os.environ["COG_RUNTIME_SCHEDULER"] = "0"  # don't fire background jobs

        # Force-reload db / app modules so the new env is picked up.
        for mod in list(sys.modules):
            if mod in {"db", "app"} or mod.startswith(("modules.", "services.")):
                del sys.modules[mod]

        from app import create_app
        from db import get_conn

        app = create_app()
        app.testing = True
        client = app.test_client()

        # --- Seed: a station + station_profile ---
        conn = get_conn(); cur = conn.cursor()
        cur.execute(
            "INSERT INTO stations (brand, name, code, station_number, state, city, address) "
            "VALUES ('consulting','Estacion Piloto','C-PIL','9001','CDMX','Iztapalapa','Av. Test 123')"
        )
        sid = cur.lastrowid
        cur.execute(
            "INSERT INTO station_profiles (station_id, brand, permit_number, legal_name, "
            "rfc, domicilio, permiso_cre, representante_legal, responsable_sasisopa, telefono) "
            "VALUES (?,?,?,?,?,?,?,?,?,?)",
            (
                sid, "consulting", "PERM-9001", "Estacion Piloto S.A. de C.V.",
                "EPI240101AAA", "Av. Test 123, CDMX",
                "PL/9001/EXP/ES/2026", "Lic. Test",
                "Ing. Pérez", "55-5555-9001",
            ),
        )
        conn.commit(); conn.close()

        # --- 1. Login ---
        print("\n[1] Login")
        resp = client.post("/api/auth/login", json={"username": "admin", "password": "admin123"})
        check("login returns 200", resp.status_code == 200, f"status={resp.status_code} body={resp.get_data(as_text=True)[:200]}")
        check("login response ok=True", resp.get_json(silent=True) and resp.get_json().get("ok") is True)

        # --- 2. Upload template ---
        print("\n[2] Upload template")
        sample = make_sample_docx_bytes()
        resp = client.post(
            "/admin/docx/templates",
            data={
                "file": (io.BytesIO(sample), "bitacora_sasisopa.docx"),
                "code": "bitacora_sasisopa",
                "name": "Bitacora SASISOPA Diaria",
                "module": "sasisopa",
                "description": "Bitacora diaria de revision SASISOPA",
            },
            content_type="multipart/form-data",
        )
        check("upload returns 200", resp.status_code == 200, f"status={resp.status_code} body={resp.get_data(as_text=True)[:300]}")
        body = resp.get_json() or {}
        check("upload ok=True", body.get("ok") is True)
        template_id = int(body.get("template_id") or 0)
        version_id = int(body.get("version_id") or 0)
        check("template_id assigned", template_id > 0)
        check("version_id assigned", version_id > 0)
        check("version label is v1.0", body.get("version_label") == "v1.0")
        detected_vars = {f["variable_name"] for f in body.get("fields", [])}
        check("RFC field detected", "RFC" in detected_vars)
        check("OBSERVACIONES field detected", "OBSERVACIONES" in detected_vars)
        check("HALLAZGOS field detected (unknown -> manual)", "HALLAZGOS" in detected_vars)
        rfc_field = next((f for f in body.get("fields", []) if f["variable_name"] == "RFC"), None)
        hallazgos_field = next((f for f in body.get("fields", []) if f["variable_name"] == "HALLAZGOS"), None)
        check("RFC classified as auto", (rfc_field or {}).get("field_kind") == "auto")
        check("HALLAZGOS classified as manual", (hallazgos_field or {}).get("field_kind") == "manual")

        # --- 3. Reject duplicate code ---
        print("\n[3] Reject duplicate code")
        resp = client.post(
            "/admin/docx/templates",
            data={
                "file": (io.BytesIO(sample), "bitacora_sasisopa.docx"),
                "code": "bitacora_sasisopa", "name": "Dup", "module": "sasisopa",
            },
            content_type="multipart/form-data",
        )
        check("duplicate upload returns 409", resp.status_code == 409)

        # --- 4. List templates ---
        print("\n[4] List templates")
        resp = client.get("/admin/docx/templates")
        check("list returns 200", resp.status_code == 200)
        body = resp.get_json() or {}
        ids = [t["id"] for t in body.get("templates", [])]
        check("template appears in list", template_id in ids)

        # --- 5. Field config: rename + mark required ---
        print("\n[5] Edit field config")
        resp = client.get(f"/admin/docx/templates/{template_id}/fields")
        check("get fields returns 200", resp.status_code == 200)
        fields = (resp.get_json() or {}).get("fields", [])
        for f in fields:
            if f["variable_name"] == "OBSERVACIONES":
                f["label"] = "Observaciones del dia"
                f["is_required"] = True
            elif f["variable_name"] == "HALLAZGOS":
                f["label"] = "Hallazgos detectados"
        resp = client.post(
            f"/admin/docx/templates/{template_id}/fields",
            json={"fields": fields},
        )
        check("save fields returns 200", resp.status_code == 200, resp.get_data(as_text=True)[:200])
        body = resp.get_json() or {}
        check("at least 2 rows updated", body.get("updated", 0) >= 2)
        obs_field = next((f for f in body.get("fields", []) if f["variable_name"] == "OBSERVACIONES"), None)
        check("OBSERVACIONES label persisted", (obs_field or {}).get("label") == "Observaciones del dia")
        check("OBSERVACIONES is_required persisted", (obs_field or {}).get("is_required") == 1)

        # --- 6. Publish ---
        print("\n[6] Publish template")
        resp = client.post(f"/admin/docx/templates/{template_id}/publish", json={"is_published": True})
        check("publish returns 200", resp.status_code == 200)
        check("is_published=True", (resp.get_json() or {}).get("is_published") is True)

        # --- 7. Generate document #1 ---
        print("\n[7] Generate document")
        resp = client.post("/admin/docx/generate", json={
            "template_id": template_id,
            "station_id": sid,
            "manual_values": {
                "OBSERVACIONES": "Todo en orden, sin incidencias.",
                "HALLAZGOS": "Ninguno.",
            },
        })
        check("generate returns 200", resp.status_code == 200, resp.get_data(as_text=True)[:300])
        gen = (resp.get_json() or {}).get("generated") or {}
        gen_id = int(gen.get("id") or 0)
        check("generated_id assigned", gen_id > 0)
        check("status starts as borrador", gen.get("status") == "borrador")
        check("docx_path populated", bool(gen.get("docx_path")))
        check("pdf_path empty (PDF backend not wired)", not gen.get("pdf_path"))

        # --- 8. List generated ---
        print("\n[8] List generated")
        resp = client.get(f"/admin/docx/generated?template_id={template_id}")
        check("list generated returns 200", resp.status_code == 200)
        ids = [g["id"] for g in (resp.get_json() or {}).get("generated", [])]
        check("generated doc appears in list", gen_id in ids)

        # --- 9. Download .docx and verify replacements ---
        print("\n[9] Download generated .docx and verify content")
        resp = client.get(f"/admin/docx/generated/{gen_id}/download")
        check("download returns 200", resp.status_code == 200)
        check(
            "content-type is .docx",
            "wordprocessingml" in (resp.headers.get("Content-Type") or "")
            or "officedocument" in (resp.headers.get("Content-Type") or ""),
            f"content-type={resp.headers.get('Content-Type')}",
        )
        content = resp.get_data()
        text = docx_text_from_bytes(content)
        check("rendered text contains station name", "Estacion Piloto" in text)
        check("rendered text contains resolved RFC", "EPI240101AAA" in text)
        check("rendered text contains permiso CRE", "PL/9001/EXP/ES/2026" in text)
        check("rendered text contains manual observation", "Todo en orden" in text)
        check("rendered text contains manual hallazgos", "Ninguno." in text)
        check("no <<RFC>> placeholder leaked", "<<RFC>>" not in text)
        check("no <<OBSERVACIONES>> placeholder leaked", "<<OBSERVACIONES>>" not in text)

        # --- 10. Approve ---
        print("\n[10] Approve")
        resp = client.post(f"/admin/docx/generated/{gen_id}/approve")
        check("approve returns 200", resp.status_code == 200)
        check("status becomes aprobado", (resp.get_json() or {}).get("status") == "aprobado")

        # --- 11. Generate doc #2 and cancel ---
        print("\n[11] Cancel another doc and verify terminal state")
        resp = client.post("/admin/docx/generate", json={
            "template_id": template_id,
            "station_id": sid,
            "manual_values": {"OBSERVACIONES": "Para cancelar.", "HALLAZGOS": "-"},
        })
        gen2_id = int(((resp.get_json() or {}).get("generated") or {}).get("id") or 0)
        check("second doc generated", gen2_id > 0)

        resp = client.post(f"/admin/docx/generated/{gen2_id}/cancel",
                            json={"reason": "Datos incorrectos"})
        check("cancel returns 200", resp.status_code == 200)
        check("status becomes cancelado", (resp.get_json() or {}).get("status") == "cancelado")

        resp = client.post(f"/admin/docx/generated/{gen2_id}/cancel",
                            json={"reason": "Otra vez"})
        check(
            "second cancel rejected (terminal state)",
            resp.status_code == 400 and (resp.get_json() or {}).get("error") == "invalid_transition",
            f"status={resp.status_code} body={resp.get_data(as_text=True)[:200]}",
        )

        # --- 12. Upload a new version (v1.1) ---
        print("\n[12] Upload new version")
        sample_v2 = make_sample_docx_bytes()  # same content for simplicity
        resp = client.post(
            f"/admin/docx/templates/{template_id}/versions",
            data={
                "file": (io.BytesIO(sample_v2), "bitacora_sasisopa_v2.docx"),
                "notes": "Sin cambios funcionales (smoke test)",
            },
            content_type="multipart/form-data",
        )
        check("new version returns 200", resp.status_code == 200, resp.get_data(as_text=True)[:300])
        body = resp.get_json() or {}
        check("new version label is v1.1", body.get("version_label") == "v1.1")
        new_version_id = int(body.get("version_id") or 0)
        check("new_version_id is different", new_version_id != version_id)

        # Verify carry-over: the OBSERVACIONES label/is_required from v1.0 must persist on v1.1.
        resp = client.get(f"/admin/docx/templates/{template_id}/fields")
        body = resp.get_json() or {}
        obs_field = next((f for f in body.get("fields", []) if f["variable_name"] == "OBSERVACIONES"), None)
        check(
            "field config carried over to new version",
            (obs_field or {}).get("label") == "Observaciones del dia"
            and (obs_field or {}).get("is_required") == 1,
            f"got: {obs_field}",
        )

    print()
    if FAILED:
        print(f"FAILED: {len(FAILED)} check(s)")
        for label in FAILED:
            print(f"  - {label}")
        return 1
    print("ALL PASS")
    return 0


if __name__ == "__main__":
    sys.exit(main())
