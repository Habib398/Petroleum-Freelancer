"""Hands-on DOCX engine demo.

Walks through the full admin flow and **leaves the generated documents on
disk** so you can open them in Word/LibreOffice and verify the variable
substitution worked. Uses a throw-away SQLite DB and uploads dir so the real
project data is never touched.

Run from the project root::

    .venv/Scripts/python.exe scripts/demo_docx.py

When it finishes you'll see something like::

    Demo files saved in:
      C:\\...\\demo_output\\
    Open the files in Word/LibreOffice to inspect the result:
      - 01_template_master.docx       (what admin uploaded)
      - 02_doc_aprobado.docx          (filled + approved document)
      - 03_doc_borrador.docx          (filled but still in 'borrador')
      - README.txt                     (what each file demonstrates)
"""

from __future__ import annotations

import io
import os
import shutil
import sys
import tempfile
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

OUTPUT_DIR = ROOT / "demo_output"


# ---------------------------------------------------------------------------
# Build a sample template that exercises auto, manual, and system variables.
# ---------------------------------------------------------------------------

def make_master_template() -> bytes:
    from docx import Document
    doc = Document()

    doc.add_heading("BITACORA DE INSPECCION SASISOPA", level=1)

    doc.add_paragraph("Estacion:           <<NOMBRE_ESTACION>>")
    doc.add_paragraph("Numero:             <<NUMERO_ESTACION>>")
    doc.add_paragraph("Razon social:       <<RAZON_SOCIAL>>")
    doc.add_paragraph("RFC:                <<RFC>>")
    doc.add_paragraph("Domicilio:          <<DOMICILIO>>")
    doc.add_paragraph("Permiso CRE:        <<PERMISO_CRE>>")
    doc.add_paragraph("Representante:      <<REPRESENTANTE_LEGAL>>")
    doc.add_paragraph("Resp. SASISOPA:     <<RESPONSABLE_SASISOPA>>")
    doc.add_paragraph("Telefono:           <<TELEFONO_ESTACION>>")
    doc.add_paragraph("")
    doc.add_paragraph("Fecha de inspeccion:    <<FECHA_HOY>>")
    doc.add_paragraph("")

    doc.add_heading("Resultado de la inspeccion", level=2)
    doc.add_paragraph("Observaciones:")
    doc.add_paragraph("<<OBSERVACIONES>>")
    doc.add_paragraph("")
    doc.add_paragraph("Hallazgos detectados:")
    doc.add_paragraph("<<HALLAZGOS>>")
    doc.add_paragraph("")
    doc.add_paragraph("Medidas correctivas:")
    doc.add_paragraph("<<MEDIDAS_CORRECTIVAS>>")
    doc.add_paragraph("")

    doc.add_paragraph("Firma del responsable: <<RESPONSABLE_SASISOPA>>")
    doc.add_paragraph("Fecha de firma:        <<FECHA_HOY>>")

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Demo
# ---------------------------------------------------------------------------

def run() -> int:
    if OUTPUT_DIR.exists():
        shutil.rmtree(OUTPUT_DIR)
    OUTPUT_DIR.mkdir(parents=True)

    tmp = Path(tempfile.mkdtemp(prefix="docx_demo_"))
    db_path = tmp / "demo.db"
    upload_dir = tmp / "uploads"
    upload_dir.mkdir(parents=True, exist_ok=True)

    os.environ["COG_DB_PATH"] = str(db_path)
    os.environ["COG_UPLOAD_DIR"] = str(upload_dir)
    os.environ["COG_CSRF"] = "0"
    os.environ["COG_SECRET"] = "demo-secret"
    os.environ["COG_ADMIN_USER"] = "admin"
    os.environ["COG_ADMIN_PASS"] = "admin123"
    os.environ["COG_RUNTIME_SCHEDULER"] = "0"

    for mod in list(sys.modules):
        if mod in {"db", "app"} or mod.startswith(("modules.", "services.")):
            del sys.modules[mod]

    from app import create_app
    from db import get_conn

    app = create_app()
    app.testing = True
    client = app.test_client()

    # 1. Login
    print("[1] Login as admin...")
    r = client.post("/api/auth/login", json={"username": "admin", "password": "admin123"})
    assert r.status_code == 200, f"login failed: {r.status_code} {r.get_data(as_text=True)}"
    print("    OK")

    # 2. Seed a station with realistic data so auto-resolution has values.
    print("[2] Seeding station with private data...")
    conn = get_conn(); cur = conn.cursor()
    cur.execute(
        "INSERT INTO stations (brand, name, code, station_number, state, city, address) "
        "VALUES ('consulting','Estacion Las Choapas','C-CHO-001','4521','Veracruz','Las Choapas','Carretera Coatza-Villahermosa km 12.5')"
    )
    sid = cur.lastrowid
    cur.execute(
        "INSERT INTO station_profiles (station_id, brand, permit_number, legal_name, "
        "rfc, domicilio, permiso_cre, representante_legal, responsable_sasisopa, telefono) "
        "VALUES (?,?,?,?,?,?,?,?,?,?)",
        (
            sid, "consulting",
            "PER/123456/2025",
            "Servicios Energeticos del Sureste S.A. de C.V.",
            "SES220315ABC",
            "Carretera Coatza-Villahermosa km 12.5, Las Choapas, Ver. CP 96980",
            "PL/12345/EXP/ES/2025",
            "Lic. Maria Lopez Hernandez",
            "Ing. Juan Perez Castillo",
            "(921) 555-4521",
        ),
    )
    conn.commit(); conn.close()
    print(f"    Station id={sid} seeded with full private data")

    # 3. Upload the master template
    print("[3] Uploading master template (.docx)...")
    master = make_master_template()
    (OUTPUT_DIR / "01_template_master.docx").write_bytes(master)

    r = client.post(
        "/admin/docx/templates",
        data={
            "file": (io.BytesIO(master), "bitacora_sasisopa.docx"),
            "code": "bitacora_sasisopa",
            "name": "Bitacora SASISOPA",
            "module": "sasisopa",
            "description": "Bitacora diaria de inspeccion SASISOPA",
        },
        content_type="multipart/form-data",
    )
    assert r.status_code == 200, r.get_data(as_text=True)
    body = r.get_json()
    template_id = body["template_id"]
    detected = body["fields"]
    print(f"    Template id={template_id}, version={body['version_label']}, {len(detected)} variables detected:")
    for f in detected:
        kind = f["field_kind"]
        marker = "[AUTO]   " if kind == "auto" else \
                 "[MANUAL] " if kind == "manual" else \
                 "[FECHA]  " if kind == "date_today" else \
                 "[IMG]    " if kind == "image" else \
                 f"[{kind:7}]"
        print(f"      {marker} <<{f['variable_name']:25}>> -> {f.get('auto_source') or '(usuario lo llenara al generar)'}")

    # 4. Publish template
    print("[4] Publishing template...")
    r = client.post(f"/admin/docx/templates/{template_id}/publish", json={"is_published": True})
    assert r.status_code == 200
    print("    OK, is_published=True")

    # 5a. Generate doc #1: with rich manual values, then APPROVE
    print("[5] Generating document #1 (with manual values + approval)...")
    r = client.post("/admin/docx/generate", json={
        "template_id": template_id,
        "station_id": sid,
        "manual_values": {
            "OBSERVACIONES": (
                "La inspeccion se realizo conforme al programa anual. "
                "Se verifico el estado de bombas, mangueras, sistema de "
                "deteccion de fugas y senalizacion."
            ),
            "HALLAZGOS": "Se detecto desgaste en una manguera de Magna (bahia 3).",
            "MEDIDAS_CORRECTIVAS": (
                "Reemplazo programado de la manguera para el dia siguiente. "
                "Bahia bloqueada con conos en lo que se realiza el cambio."
            ),
        },
    })
    assert r.status_code == 200, r.get_data(as_text=True)
    gen1 = r.get_json()["generated"]
    print(f"    Generated id={gen1['id']}, status={gen1['status']}, docx_path={gen1['docx_path']}")

    r = client.post(f"/admin/docx/generated/{gen1['id']}/approve")
    assert r.status_code == 200
    print(f"    Approved -> status={r.get_json()['status']}")

    r = client.get(f"/admin/docx/generated/{gen1['id']}/download")
    assert r.status_code == 200
    (OUTPUT_DIR / "02_doc_aprobado.docx").write_bytes(r.get_data())
    print(f"    Downloaded to demo_output/02_doc_aprobado.docx ({len(r.get_data())} bytes)")

    # 5b. Generate doc #2: minimal manual values, leave in BORRADOR
    print("[6] Generating document #2 (left in 'borrador' for comparison)...")
    r = client.post("/admin/docx/generate", json={
        "template_id": template_id,
        "station_id": sid,
        "manual_values": {
            "OBSERVACIONES": "Inspeccion rutinaria sin novedad.",
            "HALLAZGOS": "Ninguno.",
            "MEDIDAS_CORRECTIVAS": "No aplica.",
        },
    })
    gen2 = r.get_json()["generated"]
    r = client.get(f"/admin/docx/generated/{gen2['id']}/download")
    (OUTPUT_DIR / "03_doc_borrador.docx").write_bytes(r.get_data())
    print(f"    Generated id={gen2['id']}, status={gen2['status']}")
    print(f"    Downloaded to demo_output/03_doc_borrador.docx")

    # 6. Final: list everything
    print("[7] Final state:")
    r = client.get("/admin/docx/templates")
    for t in r.get_json()["templates"]:
        print(f"    Template: {t['code']} '{t['name']}' v{t.get('current_version_label')} fields={t.get('field_count')} published={t.get('is_published')}")
    r = client.get("/admin/docx/generated")
    for g in r.get_json()["generated"]:
        print(f"    Generated: id={g['id']} station={g.get('station_name')} status={g['status']} created_at={g['created_at']}")

    # 7. README so anyone receiving the demo folder understands what it shows
    readme = OUTPUT_DIR / "README.txt"
    readme.write_text(
        "Demo del motor de plantillas DOCX\n"
        "==================================\n\n"
        "Esta carpeta contiene 3 archivos generados por el script demo_docx.py.\n"
        "Abrelos en Word o LibreOffice para verificar el funcionamiento.\n\n"
        "01_template_master.docx\n"
        "  Plantilla original que el admin subio. Contiene placeholders <<...>>\n"
        "  como <<RFC>>, <<NOMBRE_ESTACION>>, <<OBSERVACIONES>>, etc.\n\n"
        "02_doc_aprobado.docx\n"
        "  Documento generado a partir de la plantilla, con datos de la estacion\n"
        "  ya rellenados automaticamente (RFC, razon social, permiso CRE...) y\n"
        "  observaciones/hallazgos/medidas correctivas escritas por el admin.\n"
        "  Estado: aprobado.\n\n"
        "03_doc_borrador.docx\n"
        "  Otro documento generado para la misma estacion pero con observaciones\n"
        "  minimas. Estado: borrador.\n\n"
        "Verifica que en los archivos 02 y 03:\n"
        "  - <<RFC>> aparece como SES220315ABC\n"
        "  - <<NOMBRE_ESTACION>> aparece como Estacion Las Choapas\n"
        "  - <<PERMISO_CRE>> aparece como PL/12345/EXP/ES/2025\n"
        "  - <<FECHA_HOY>> aparece con la fecha actual\n"
        "  - Las observaciones/hallazgos son las que se escribieron al generar.\n",
        encoding="utf-8",
    )

    # Cleanup the temp uploads/db; keep only the curated demo_output/.
    shutil.rmtree(tmp, ignore_errors=True)

    print()
    print("=" * 60)
    print("Demo files saved in:")
    print(f"  {OUTPUT_DIR}")
    print()
    print("Open the files in Word / LibreOffice to inspect:")
    print("  - 01_template_master.docx   (plantilla original con <<...>>)")
    print("  - 02_doc_aprobado.docx      (rellenado + aprobado)")
    print("  - 03_doc_borrador.docx      (rellenado + en borrador)")
    print("  - README.txt                (que verificar en cada archivo)")
    print("=" * 60)
    return 0


if __name__ == "__main__":
    sys.exit(run())
