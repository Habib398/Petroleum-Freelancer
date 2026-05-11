"""End-to-end smoke test for the DOCX template engine.

Creates a sample ``.docx`` template programmatically (so the test is
self-contained), parses it, renders it with sample values, then re-opens the
rendered file to verify every placeholder was replaced as expected.

Run from the project root::

    .venv/Scripts/python.exe scripts/smoke_docx_engine.py

Exits with status 0 on success, 1 on any failure.
"""

from __future__ import annotations

import sys
import tempfile
from pathlib import Path

# Make the project root importable when running this script directly.
ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

from docx import Document  # noqa: E402

from services.docx_engine import parse_template, render_docx  # noqa: E402
from services.docx_variables import (  # noqa: E402
    classify,
    list_known,
    merge_with_manual,
    resolve_auto_values,
)


# ---------------------------------------------------------------------------
# Test helpers
# ---------------------------------------------------------------------------

FAILED: list[str] = []


def check(label: str, condition: bool, detail: str = "") -> None:
    if condition:
        print(f"  OK  {label}")
    else:
        FAILED.append(label)
        print(f"  FAIL  {label}")
        if detail:
            print(f"        {detail}")


def make_sample_template(path: Path) -> None:
    """Create a small but representative .docx template on disk."""
    doc = Document()
    doc.add_heading("BITACORA SASISOPA", level=1)

    doc.add_paragraph("Estación: <<NOMBRE_ESTACION>>")
    doc.add_paragraph("Número: <<NUMERO_ESTACION>>")
    doc.add_paragraph("Razón social: <<RAZON_SOCIAL>>")
    doc.add_paragraph("RFC: <<RFC>>")
    doc.add_paragraph("Permiso CRE: <<PERMISO_CRE>>")
    doc.add_paragraph("Fecha: <<FECHA_HOY>>")

    doc.add_paragraph("Observaciones:")
    doc.add_paragraph("<<OBSERVACIONES>>")

    # A duplicate placeholder (should not double-count in parse)
    doc.add_paragraph("Confirmamos RFC <<RFC>> al cierre.")

    # An unknown variable that admin will need to classify manually
    doc.add_paragraph("Hallazgos: <<HALLAZGOS>>")

    # A variable inside a table cell
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "Responsable SASISOPA"
    table.rows[0].cells[1].text = "<<RESPONSABLE_SASISOPA>>"
    table.rows[1].cells[0].text = "Teléfono"
    table.rows[1].cells[1].text = "<<TELEFONO_ESTACION>>"

    doc.save(str(path))


def read_full_text(path: Path) -> str:
    """Concatenate every paragraph (body + tables) for assertion convenience."""
    doc = Document(str(path))
    chunks: list[str] = []
    for p in doc.paragraphs:
        chunks.append(p.text)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    chunks.append(p.text)
    return "\n".join(chunks)


# ---------------------------------------------------------------------------
# Test cases
# ---------------------------------------------------------------------------

def test_catalog() -> None:
    print("\n[1] Catalog sanity")
    catalog = list_known()
    check("catalog non-empty", len(catalog) > 0, f"got {len(catalog)} entries")
    names = {entry["variable"] for entry in catalog}
    for must_have in ("RFC", "NOMBRE_ESTACION", "PERMISO_CRE", "LOGO_ESTACION", "FECHA_HOY"):
        check(f"catalog contains <<{must_have}>>", must_have in names)

    check("classify(<<RFC>>) == auto", classify("RFC") == "auto")
    check("classify(<<LOGO_ESTACION>>) == image", classify("LOGO_ESTACION") == "image")
    check("classify(<<FECHA_HOY>>) == date_today", classify("FECHA_HOY") == "date_today")
    check("classify(<<HALLAZGOS>>) == manual", classify("HALLAZGOS") == "manual")


def test_parser(tmpdir: Path) -> list[dict]:
    print("\n[2] Parser")
    tpl = tmpdir / "bitacora_sample.docx"
    make_sample_template(tpl)

    detected = parse_template(tpl)
    by_name = {item["variable"]: item for item in detected}

    expected = {
        "NOMBRE_ESTACION", "NUMERO_ESTACION", "RAZON_SOCIAL", "RFC",
        "PERMISO_CRE", "FECHA_HOY", "OBSERVACIONES", "HALLAZGOS",
        "RESPONSABLE_SASISOPA", "TELEFONO_ESTACION",
    }
    check(
        "all placeholders detected",
        expected.issubset(by_name.keys()),
        f"missing: {expected - by_name.keys()}",
    )
    check("RFC has auto kind", by_name.get("RFC", {}).get("kind") == "auto")
    check(
        "FECHA_HOY has date_today kind",
        by_name.get("FECHA_HOY", {}).get("kind") == "date_today",
    )
    check(
        "OBSERVACIONES has manual kind",
        by_name.get("OBSERVACIONES", {}).get("kind") == "manual",
    )
    check(
        "HALLAZGOS (unknown) classified as manual",
        by_name.get("HALLAZGOS", {}).get("kind") == "manual",
    )
    check(
        "RFC occurrences == 2 (counted both uses)",
        by_name.get("RFC", {}).get("occurrences") == 2,
        f"got {by_name.get('RFC', {}).get('occurrences')}",
    )
    check(
        "RESPONSABLE_SASISOPA detected inside table cell",
        "RESPONSABLE_SASISOPA" in by_name,
    )
    return detected


def test_render(tmpdir: Path) -> None:
    print("\n[3] Render with sample values")
    tpl = tmpdir / "bitacora_sample.docx"  # already created by parser test
    out = tmpdir / "bitacora_filled.docx"

    values = {
        "NOMBRE_ESTACION": "Las Choapas",
        "NUMERO_ESTACION": "1234",
        "RAZON_SOCIAL": "Servicios HME S.A. de C.V.",
        "RFC": "SHM240101ABC",
        "PERMISO_CRE": "PL/12345/EXP/ES/2026",
        "FECHA_HOY": "2026-05-08",
        "OBSERVACIONES": "Todo en orden, sin incidencias.",
        "RESPONSABLE_SASISOPA": "Juan Pérez",
        "TELEFONO_ESTACION": "921-555-0000",
        # Note: HALLAZGOS deliberately omitted -> should remain as <<HALLAZGOS>>
    }
    render_docx(tpl, out, values=values)

    check("rendered file exists", out.exists())

    text = read_full_text(out)

    for var, expected_value in values.items():
        check(
            f"<<{var}>> replaced with {expected_value!r}",
            expected_value in text,
            f"actual snippet: {text[:300]!r}",
        )
        check(f"<<{var}>> placeholder no longer present", f"<<{var}>>" not in text)

    check(
        "omitted <<HALLAZGOS>> placeholder preserved (so admin notices)",
        "<<HALLAZGOS>>" in text,
    )


def test_render_with_db_resolution(tmpdir: Path) -> None:
    """Verify that resolve_auto_values + render produces a real document.

    Uses a temporary SQLite DB pointed at by the COG_DB env var so the real
    project DB is never touched.
    """
    print("\n[4] Render using auto-resolved values from DB")
    import os, sqlite3

    db_path = tmpdir / "smoke.db"
    os.environ["COG_DB_PATH"] = str(db_path)

    # Force a fresh import of db with the env override so it picks up our path.
    for mod in list(sys.modules):
        if mod == "db" or mod.startswith("db."):
            del sys.modules[mod]
    from db import init_db, get_conn  # noqa: WPS433  (re-import on purpose)

    init_db()
    conn = get_conn()
    cur = conn.cursor()
    cur.execute(
        "INSERT INTO stations (brand, name, code, station_number, state, city, address) "
        "VALUES ('petroleum', 'Las Choapas', 'P-001', '1234', 'Veracruz', 'Las Choapas', 'Carretera 145 km 12')"
    )
    sid = cur.lastrowid
    cur.execute(
        "INSERT INTO station_profiles (station_id, brand, permit_number, legal_name, "
        "rfc, domicilio, permiso_cre, representante_legal, responsable_sasisopa, telefono) "
        "VALUES (?, 'petroleum', ?, ?, ?, ?, ?, ?, ?, ?)",
        (
            sid, "PERM-001", "Servicios HME S.A. de C.V.",
            "SHM240101ABC", "Carretera 145 km 12, Las Choapas, Ver.",
            "PL/12345/EXP/ES/2026", "Lic. María López",
            "Juan Pérez", "921-555-0000",
        ),
    )
    conn.commit()

    from services.docx_variables import resolve_auto_values  # local re-import after env override

    auto = resolve_auto_values(sid, conn=conn)
    conn.close()

    check("auto[NOMBRE_ESTACION] populated", auto.get("NOMBRE_ESTACION") == "Las Choapas")
    check("auto[RFC] populated", auto.get("RFC") == "SHM240101ABC")
    check("auto[PERMISO_CRE] populated", auto.get("PERMISO_CRE") == "PL/12345/EXP/ES/2026")
    check("auto[FECHA_HOY] populated (system value)", isinstance(auto.get("FECHA_HOY"), str))

    merged = merge_with_manual(auto, {"OBSERVACIONES": "Bitácora de prueba."})
    check(
        "manual override merged",
        merged.get("OBSERVACIONES") == "Bitácora de prueba.",
    )
    check(
        "auto value preserved through merge",
        merged.get("RFC") == "SHM240101ABC",
    )

    tpl = tmpdir / "bitacora_sample.docx"
    out = tmpdir / "bitacora_resolved.docx"
    render_docx(tpl, out, values=merged)
    text = read_full_text(out)
    check("rendered DOCX contains resolved RFC", "SHM240101ABC" in text)
    check("rendered DOCX contains manual observation", "Bitácora de prueba." in text)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main() -> int:
    print("DOCX engine smoke test")
    print("======================")
    with tempfile.TemporaryDirectory(prefix="docx_smoke_") as tmp:
        tmpdir = Path(tmp)
        test_catalog()
        test_parser(tmpdir)
        test_render(tmpdir)
        test_render_with_db_resolution(tmpdir)

    print()
    if FAILED:
        print(f"FAILED: {len(FAILED)} check(s)")
        for label in FAILED:
            print(f"  - {label}")
        return 1
    print("ALL PASS")
    return 0


if __name__ == "__main__":
    sys.exit(main())
