"""DOCX template engine: parse ``<<VARIABLE>>`` placeholders and render filled documents.

The engine has two responsibilities:

1. **Parsing** — open a ``.docx`` admin uploaded as a template, walk every text
   container (body paragraphs, table cells, headers, footers) and return the
   set of variables that appear inside ``<<…>>`` markers.

2. **Rendering** — given the same template plus a values dictionary, produce a
   new ``.docx`` where every placeholder has been replaced by the actual
   value. Image placeholders (``<<LOGO_EMPRESA>>`` etc.) are replaced by
   inserting the picture; text/date placeholders by string substitution.

Run-splitting in Word
---------------------
Microsoft Word splits paragraph text into "runs" (each run is a stretch of
text with uniform formatting). A placeholder typed in one go can still end up
straddling several runs after Word's autoformatting (e.g. autocorrect, spell
check, copy/paste). To handle this, the renderer joins all runs of a paragraph
into a single text buffer, performs the replacements, and rewrites the result
into the first run while clearing the rest. This preserves paragraph-level
formatting (alignment, indentation, style) but loses run-level formatting
*within* paragraphs that contain placeholders. Templates that need bold-only-
on-the-value can work around this by placing the variable on its own line.
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Any, Iterable

from docx import Document
from docx.document import Document as _DocxDocument
from docx.shared import Mm
from docx.text.paragraph import Paragraph

from services.docx_variables import (
    KNOWN_VARIABLES,
    auto_source_for,
    canonical,
    classify,
    label_for,
)


# ---------------------------------------------------------------------------
# Constants
# ---------------------------------------------------------------------------

# Match ``<<VAR_NAME>>``. Variable names accept letters, digits and underscores.
# Whitespace inside the brackets is tolerated (``<< RFC >>`` is valid).
VAR_PATTERN = re.compile(r"<<\s*([A-Za-z0-9_]+)\s*>>")

# Default image width (mm) when inserting a logo. Tweakable per-field later.
DEFAULT_IMAGE_WIDTH_MM = 30


# ---------------------------------------------------------------------------
# Parsing
# ---------------------------------------------------------------------------

def _iter_paragraphs(doc: _DocxDocument) -> Iterable[Paragraph]:
    """Yield every paragraph in the document tree.

    Covers body paragraphs, table cells (recursively for nested tables),
    and headers/footers of every section.
    """
    # Body
    for p in doc.paragraphs:
        yield p

    # Body tables (and any nested tables inside them)
    def _walk_table(tbl):
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p
                for nested in cell.tables:
                    yield from _walk_table(nested)

    for tbl in doc.tables:
        yield from _walk_table(tbl)

    # Headers and footers (per section). Each can also have tables.
    for section in doc.sections:
        for hf in (section.header, section.footer,
                   section.first_page_header, section.first_page_footer,
                   section.even_page_header, section.even_page_footer):
            if hf is None:
                continue
            for p in hf.paragraphs:
                yield p
            for tbl in hf.tables:
                yield from _walk_table(tbl)


def _paragraph_text(paragraph: Paragraph) -> str:
    """Return the joined text of a paragraph (concatenates all its runs)."""
    return "".join(run.text or "" for run in paragraph.runs)


def parse_template(docx_path: str | Path) -> list[dict]:
    """Open a ``.docx`` template and return its detected variables.

    Returns a list of dicts (one per unique variable, in order of first
    appearance):

    .. code-block:: python

        [
            {
                "variable": "NOMBRE_ESTACION",
                "kind": "auto",                 # auto | manual | image | date_today
                "auto_source": "station.name",  # None for manual
                "label": "Nombre estacion",
                "occurrences": 3,               # how many times it appears
            },
            ...
        ]

    The classification is the *suggested* one based on the catalog. Admin can
    override it later via the field-config UI.
    """
    docx_path = Path(docx_path)
    if not docx_path.exists():
        raise FileNotFoundError(f"template_not_found: {docx_path}")

    doc = Document(str(docx_path))

    seen: dict[str, dict] = {}
    for paragraph in _iter_paragraphs(doc):
        text = _paragraph_text(paragraph)
        if not text or "<<" not in text:
            continue
        for match in VAR_PATTERN.finditer(text):
            cname = canonical(match.group(1))
            if not cname:
                continue
            if cname in seen:
                seen[cname]["occurrences"] += 1
                continue
            seen[cname] = {
                "variable": cname,
                "kind": classify(cname),
                "auto_source": auto_source_for(cname),
                "label": label_for(cname),
                "occurrences": 1,
            }

    return list(seen.values())


def list_unique_variables(docx_path: str | Path) -> list[str]:
    """Convenience: just the canonical variable names found in a template."""
    return [item["variable"] for item in parse_template(docx_path)]


# ---------------------------------------------------------------------------
# Rendering
# ---------------------------------------------------------------------------

def _format_value(raw: Any) -> str:
    """Normalize a value for text replacement.

    ``None`` becomes an empty string. Other types are coerced via ``str()``.
    Whitespace is preserved so admins can format multi-line observations.
    """
    if raw is None:
        return ""
    if isinstance(raw, bool):
        return "Sí" if raw else "No"
    return str(raw)


def _replace_in_paragraph_text(paragraph: Paragraph, values: dict[str, Any]) -> bool:
    """Substitute every ``<<VAR>>`` in a paragraph using its joined text.

    Strategy: read the full paragraph text, perform substitutions on it, then
    write the result back into the first run while clearing the others. This
    is the only reliable way to handle placeholders that Word has split across
    multiple runs.

    Returns ``True`` when at least one replacement happened (signals the
    caller it may want to do post-processing such as image insertion).
    """
    text = _paragraph_text(paragraph)
    if "<<" not in text:
        return False

    def _sub(match: re.Match) -> str:
        cname = canonical(match.group(1))
        if cname not in values:
            # Unknown / not-provided variable: leave the placeholder intact so
            # admin can spot it in the output rather than silently lose it.
            return match.group(0)
        return _format_value(values[cname])

    new_text = VAR_PATTERN.sub(_sub, text)
    if new_text == text:
        return False

    if not paragraph.runs:
        paragraph.add_run(new_text)
        return True

    # Write everything into the first run; wipe the rest. The first run keeps
    # its formatting (bold/italic/font) so the substituted text inherits it.
    paragraph.runs[0].text = new_text
    for run in paragraph.runs[1:]:
        run.text = ""
    return True


def _insert_image_for_variable(paragraph: Paragraph, variable: str, image_path: str | Path) -> bool:
    """Replace a paragraph that contains ONLY ``<<VAR>>`` with an image.

    To keep things predictable in this first iteration, image variables are
    only honored when the placeholder is the **sole** content of its
    paragraph. This is the convention the documentation snippet shows
    (``<<LOGO_EMPRESA>>`` on its own line). Mixing an image variable with
    surrounding text in the same paragraph will fall through to text
    substitution (with the image path as a string), which is rarely useful
    but at least never silently drops the placeholder.

    Returns ``True`` on success, ``False`` if the paragraph isn't a sole-
    placeholder line or the image file is missing.
    """
    image_path = Path(image_path) if image_path else None
    if image_path is None or not image_path.exists():
        return False

    text = _paragraph_text(paragraph).strip()
    expected = f"<<{variable}>>".upper()
    if text.upper() != expected:
        return False

    # Clear runs and add the picture in the first run.
    for run in paragraph.runs:
        run.text = ""
    if not paragraph.runs:
        run = paragraph.add_run()
    else:
        run = paragraph.runs[0]
    try:
        run.add_picture(str(image_path), width=Mm(DEFAULT_IMAGE_WIDTH_MM))
    except Exception:
        return False
    return True


def render_docx(
    template_path: str | Path,
    output_path: str | Path,
    *,
    values: dict[str, Any],
    image_values: dict[str, str | Path] | None = None,
) -> Path:
    """Render a ``.docx`` template by replacing every placeholder.

    Parameters
    ----------
    template_path:
        Path to the source ``.docx`` (admin's master template).
    output_path:
        Where to write the rendered file. Parent directories are created.
    values:
        Mapping of canonical variable names to their text values. Keys are
        normalized (so ``"rfc"`` and ``"RFC"`` are equivalent). Missing keys
        leave the placeholder visible in the output.
    image_values:
        Optional mapping of canonical variable names to image file paths. Only
        applied for paragraphs whose entire text is a single placeholder.

    Returns
    -------
    Path
        The output path (same value as ``output_path`` for convenience).
    """
    template_path = Path(template_path)
    output_path = Path(output_path)
    if not template_path.exists():
        raise FileNotFoundError(f"template_not_found: {template_path}")

    # Normalize keys (callers may pass mixed-case names).
    norm_values: dict[str, Any] = {canonical(k): v for k, v in (values or {}).items()}
    norm_images: dict[str, Path] = {
        canonical(k): Path(v) for k, v in (image_values or {}).items() if v
    }

    doc = Document(str(template_path))

    for paragraph in _iter_paragraphs(doc):
        text = _paragraph_text(paragraph)
        if "<<" not in text:
            continue

        # Image-only paragraphs go through the picture insertion path first;
        # if that path declines (paragraph has surrounding text or file
        # missing), fall back to text substitution.
        handled_as_image = False
        for match in VAR_PATTERN.finditer(text):
            cname = canonical(match.group(1))
            if cname in norm_images and classify(cname) == "image":
                if _insert_image_for_variable(paragraph, cname, norm_images[cname]):
                    handled_as_image = True
                    break
        if handled_as_image:
            continue

        _replace_in_paragraph_text(paragraph, norm_values)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    return output_path
